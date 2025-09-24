import io
import os
import pyodbc
from docx import Document
from azure.storage.blob import BlobServiceClient
from azure.identity import DefaultAzureCredential


def get_blob_client():
    """Get blob service client with Azure Identity or connection string fallback"""
    try:
        # Try Azure Identity first (recommended for production)
        credential = DefaultAzureCredential()
        blob_cs = os.environ["BLOB_CONNECTION_STRING"]
        account_name = blob_cs.split("AccountName=")[1].split(";")[0]
        storage_url = f"https://{account_name}.blob.core.windows.net"
        return BlobServiceClient(account_url=storage_url, credential=credential)
    except Exception:
        # Fallback to connection string
        return BlobServiceClient.from_connection_string(os.environ["BLOB_CONNECTION_STRING"])


def fetch_paystub_rows(cnxn, employee_number, date_from, date_to):
    """Fetch paystub data using the stored procedure"""
    with cnxn.cursor() as cur:
        cur.execute("""
            EXEC dbo.sp_GetPaystubForRange @EmployeeNumber=?, @From=?, @To=?
        """, (employee_number, date_from, date_to))
        cols = [c[0] for c in cur.description]
        return [dict(zip(cols, r)) for r in cur.fetchall()]


def render_payslip_docx(rows):
    """Generate a paystub DOCX document from the database rows"""
    if not rows: 
        raise ValueError("No pay data in range.")
    
    employee = rows[0]["FullName"]
    empnum = rows[0]["EmployeeNumber"]

    # Create new document
    doc = Document()
    doc.add_heading("Paystub", level=1)
    doc.add_paragraph(f"Employee Name: {employee}")
    doc.add_paragraph(f"Employee Number: {empnum}")
    doc.add_paragraph(f"Pay Period: {rows[0]['PeriodStart']} to {rows[-1]['PeriodEnd']}")
    doc.add_paragraph("")

    # Earnings section
    doc.add_paragraph("EARNINGS")
    earnings_table = doc.add_table(rows=1, cols=3)
    earnings_table.rows[0].cells[0].text = "Description"
    earnings_table.rows[0].cells[1].text = "Amount"
    earnings_table.rows[0].cells[2].text = "Notes"

    gross_total = 0
    net_total = 0
    cpp_total = 0
    ei_total = 0
    
    for r in rows:
        row = earnings_table.add_row().cells
        row[0].text = f"Gross {r['PeriodStart']}–{r['PeriodEnd']}"
        row[1].text = f"{float(r['GrossAmount']):.2f}"
        row[2].text = ""
        gross_total += float(r['GrossAmount'])
    
    # Total row
    row = earnings_table.add_row().cells
    row[0].text = "Gross Total"
    row[1].text = f"{gross_total:.2f}"
    row[2].text = ""

    # Deductions section
    doc.add_paragraph("")
    doc.add_paragraph("DEDUCTIONS")
    deductions_table = doc.add_table(rows=1, cols=2)
    deductions_table.rows[0].cells[0].text = "Deduction"
    deductions_table.rows[0].cells[1].text = "Amount"
    
    for r in rows:
        cpp_total += float(r['CPP'] or 0)
        ei_total += float(r['EI'] or 0)
    
    cpp_row = deductions_table.add_row().cells
    cpp_row[0].text = "CPP"
    cpp_row[1].text = f"{cpp_total:.2f}"
    
    ei_row = deductions_table.add_row().cells
    ei_row[0].text = "EI"
    ei_row[1].text = f"{ei_total:.2f}"

    # Calculate net total
    for r in rows:
        net_total += float(r['NetAmount'])
    
    doc.add_paragraph("")
    doc.add_paragraph(f"NET PAY (sum of periods): {net_total:.2f}")

    # Save to bytes buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def upload_bytes_to_blob(container, blob_name, data_bytes):
    """Upload bytes data to Azure Blob Storage using Azure Identity"""
    blob_client = get_blob_client()
    blob_client.get_container_client(container).upload_blob(
        name=blob_name, 
        data=data_bytes, 
        overwrite=True
    )
    print(f"Uploaded {blob_name} to {container} container")
